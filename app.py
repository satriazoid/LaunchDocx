import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import re

# --- FUNGSI AUTO-TOC ---
def add_table_of_contents(paragraph):
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar)
    run._element.append(instrText)
    run._element.append(fldChar2)
    run._element.append(fldChar3)

# --- FUNGSI CLEANSING AI ---
def clean_ai_content(text):
    if not text: return ""
    text = text.replace("**", "").replace("__", "").replace("*", "").replace("_", "")
    text = re.sub(r'#+\s?', '', text)
    text = text.replace("--", "").replace("—", "").replace("–", "")
    return text.strip()

# --- FUNGSI STYLING AKADEMIK ---
def apply_academic_style(paragraph, size, bold=False, is_heading=False):
    paragraph.paragraph_format.line_spacing = 1.5
    if not is_heading:
        paragraph.paragraph_format.first_line_indent = Cm(1.0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)

def process_flexible_content(doc, raw_text):
    lines = raw_text.split('\n')
    for line in lines:
        clean_line = clean_ai_content(line)
        if not clean_line: continue
        sub_match = re.match(r'^(\d+(\.\d+)+)\s+(.*)', clean_line)
        if sub_match:
            level_str = sub_match.group(1)
            title_str = f"{level_str} {sub_match.group(3)}"
            level = 2 if level_str.count('.') == 1 else 3
            h = doc.add_heading(title_str, level=level)
            apply_academic_style(h, 12, bold=True, is_heading=True)
            h.paragraph_format.keep_with_next = True
        else:
            p = doc.add_paragraph(clean_line)
            apply_academic_style(p, 12)

# --- ENGINE GENERATOR ---
def generate_universal_doc(data, config):
    doc = Document()
    
    # 1. COVER
    cp_kat = doc.add_paragraph()
    cp_kat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp_kat.add_run(data['kategori'].upper() + "\n\n")
    apply_academic_style(cp_kat, 14, bold=True, is_heading=True)

    cp_judul = doc.add_paragraph()
    cp_judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp_judul.add_run(data['judul'].upper() + "\n\n\n")
    apply_academic_style(cp_judul, 16, bold=True, is_heading=True)
    
    cp_label = doc.add_paragraph()
    cp_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp_label.add_run("Disusun Oleh:")
    apply_academic_style(cp_label, 12, bold=True, is_heading=True)

    list_nama = data['nama'].split('\n')
    list_id = data['id_mhs'].split('\n')
    for i in range(len(list_nama)):
        if list_nama[i].strip():
            p = doc.add_paragraph(f"{list_nama[i].strip()} ({list_id[i].strip() if i < len(list_id) else '-'})")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            apply_academic_style(p, 12, is_heading=True)

    doc.add_paragraph("\n\n")
    cp_inst = doc.add_paragraph()
    cp_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp_inst.add_run(f"{data['prodi'].upper()}\n{data['institusi'].upper()}\n{data['tahun']}")
    apply_academic_style(cp_inst, 14, bold=True, is_heading=True)
    doc.add_page_break()

    # 2. HALAMAN FORMAL (MODULAR)
    formal_configs = [
        ('use_pengantar', "KATA PENGANTAR"),
        ('use_daftar_isi', "DAFTAR ISI"),
        ('use_daftar_tabel', "DAFTAR TABEL"),
        ('use_daftar_gambar', "DAFTAR GAMBAR")
    ]

    for config_key, title in formal_configs:
        if config[config_key]:
            h = doc.add_heading(title, level=1)
            apply_academic_style(h, 14, bold=True, is_heading=True)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if title == "DAFTAR ISI":
                p_toc = doc.add_paragraph()
                add_table_of_contents(p_toc)
            elif data.get(title):
                process_flexible_content(doc, data[title])
            doc.add_page_break()

    # 3. BAB DINAMIS
    for i in range(config['num_chapters']):
        romawi = ["I", "II", "III", "IV", "V", "VI", "VII"][i]
        judul_bab = data[f"judul_bab_{i+1}"].upper()
        h_bab = doc.add_heading(f"BAB {romawi}\n{judul_bab}", level=1)
        apply_academic_style(h_bab, 14, bold=True, is_heading=True)
        h_bab.alignment = WD_ALIGN_PARAGRAPH.CENTER
        process_flexible_content(doc, data[f"isi_bab_{i+1}"])
        doc.add_page_break()

    # 4. DAFTAR PUSTAKA
    if config['use_pustaka']:
        h_ref = doc.add_heading("DAFTAR PUSTAKA", level=1)
        apply_academic_style(h_ref, 14, bold=True, is_heading=True)
        h_ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if data.get("DAFTAR PUSTAKA"):
            process_flexible_content(doc, data["DAFTAR PUSTAKA"])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- UI STREAMLIT ---
st.set_page_config(page_title="Universal Doc Engine", layout="wide")
st.title("Universal Academic Document Engine")

with st.sidebar:
    st.header("Identitas Institusi")
    institusi = st.text_input("Nama Institusi", "Universitas Indonesia")
    prodi = st.text_input("Jurusan/Prodi", "Teknik Informatika")
    tahun = st.text_input("Tahun/Lokasi", "Jakarta, 2026")
    
    st.divider()
    st.header("Konfigurasi Dokumen")
    kategori = st.selectbox("Kategori Dokumen", ["Makalah", "Laporan Praktikum", "Tugas Akhir", "Skripsi", "Custom"])
    num_chapters = st.number_input("Jumlah Bab", 1, 7, 3)
    
    st.subheader("Opsi Halaman")
    use_pengantar = st.checkbox("Kata Pengantar", True)
    use_daftar_isi = st.checkbox("Daftar Isi Otomatis", True)
    use_daftar_tabel = st.checkbox("Daftar Tabel", False)
    use_daftar_gambar = st.checkbox("Daftar Gambar", False)
    use_pustaka = st.checkbox("Daftar Pustaka", True)
    
    st.divider()
    st.header("Identitas Penulis")
    judul = st.text_input("Judul Tugas")
    nama_all = st.text_area("Nama Anggota (Per baris)")
    id_all = st.text_area("NIM/ID Siswa (Per baris)")

# MEMBANGUN TABS
tab_titles = ["Halaman Depan"] 
for i in range(num_chapters): tab_titles.append(f"Bab {i+1}")
if use_pustaka: tab_titles.append("Daftar Pustaka")

tabs = st.tabs(tab_titles)
data_in = {
    'institusi': institusi, 'prodi': prodi, 'tahun': tahun, 
    'kategori': kategori, 'judul': judul, 'nama': nama_all, 'id_mhs': id_all
}

with tabs[0]:
    if use_pengantar:
        data_in["KATA PENGANTAR"] = st.text_area("Isi Kata Pengantar", height=200)
    if use_daftar_tabel:
        data_in["DAFTAR TABEL"] = st.text_area("Isi Daftar Tabel (Opsional)", height=150)
    if use_daftar_gambar:
        data_in["DAFTAR GAMBAR"] = st.text_area("Isi Daftar Gambar (Opsional)", height=150)

for i in range(num_chapters):
    with tabs[i+1]:
        data_in[f"judul_bab_{i+1}"] = st.text_input(f"Judul BAB {i+1}", value="Pendahuluan", key=f"t_{i}")
        data_in[f"isi_bab_{i+1}"] = st.text_area(f"Konten", height=400, key=f"c_{i}")

if use_pustaka:
    with tabs[-1]:
        data_in["DAFTAR PUSTAKA"] = st.text_area("Referensi", height=400)

st.divider()
col_pre, col_gen = st.columns(2)

with col_pre:
    if st.button("Preview Konten", use_container_width=True):
        st.subheader("Pratinjau Teks Bersih")
        with st.container(border=True):
            if use_pengantar:
                with st.expander("KATA PENGANTAR"):
                    st.write(clean_ai_content(data_in.get("KATA PENGANTAR", "")))
            for i in range(num_chapters):
                with st.expander(f"BAB {i+1}: {data_in[f'judul_bab_{i+1}'].upper()}"):
                    st.write(clean_ai_content(data_in[f"isi_bab_{i+1}"]))
            if use_pustaka:
                with st.expander("DAFTAR PUSTAKA"):
                    st.write(clean_ai_content(data_in.get("DAFTAR PUSTAKA", "")))

with col_gen:
    if st.button("Generate dan Download .docx", use_container_width=True):
        if not (judul and nama_all):
            st.error("Lengkapi Judul dan Nama Penulis!")
        else:
            conf = {
                'num_chapters': num_chapters, 
                'use_pengantar': use_pengantar, 
                'use_daftar_isi': use_daftar_isi,
                'use_daftar_tabel': use_daftar_tabel,
                'use_daftar_gambar': use_daftar_gambar,
                'use_pustaka': use_pustaka
            }
            file = generate_universal_doc(data_in, conf)
            st.download_button("Unduh Dokumen", file, f"Dokumen_{kategori}.docx")